#!/usr/bin/env python3
"""
pdf_to_word_tracked_changes.py

Converts Colorado General Assembly bill PDFs to Word documents with tracked changes.

Uses character-level PDF extraction to detect:
  - Struck-through text (tagged 'ocr_q' in the PDF) → Word tracked deletion
  - ALL CAPS text → Word tracked insertion (new law)
  - Normal mixed-case text → unchanged

Author attribution for tracked changes: "Colorado General Assembly"

Usage:
    python pdf_to_word.py input.pdf output.docx

Requirements:
    pip install pdfplumber python-docx lxml
"""

import re
import sys
import argparse
from pathlib import Path
from copy import deepcopy
from collections import defaultdict

import pdfplumber
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches


# ── Constants ────────────────────────────────────────────────────────────────

AUTHOR = "Colorado General Assembly"
DATE = "2021-01-01T00:00:00Z"

DASH_CHARS = set("-\u2013\u2014")

# Noise lines to skip (headers, footers, signature block, etc.)
NOISE_RE = re.compile(
    "|".join([
        r"^PAGE \d+[-\u2013\u2014](?:HOUSE|SENATE) BILL",
        r"^Capital letters or bold",
        r"^through words or numbers",
        r"^the act\.",
        r"^SPEAKER OF THE HOUSE",
        r"^OF REPRESENTATIVES",
        r"^CHIEF CLERK",
        r"^PRESIDENT OF",
        r"^THE SENATE$",
        r"^SECRETARY OF",
        r"^APPROVED\s",
        r"^GOVERNOR OF THE STATE",
        r"^Jared S\.",
        r"^Alec Garnett",
        r"^Robin Jones",
        r"^Cindi L\.",
        r"^Leroy M\.",
        r"^Le M\. Garcia",
    ])
)

# Abbreviations like U.S.C. — always normal, not insertions
ABBREV_RE = re.compile(r"^([A-Z]\.){2,}[A-Z]?$")

# Legal citations like 24-34-301 — always normal
CITATION_RE = re.compile(r"^\d{1,3}-\d{1,5}-\d{1,5}")


# ── Tracked-change ID counter ────────────────────────────────────────────────

_change_id = 0


def _next_id() -> str:
    global _change_id
    _change_id += 1
    return str(_change_id)


# ── XML helpers ───────────────────────────────────────────────────────────────

def _base_rpr() -> OxmlElement:
    """Build a base <w:rPr> with Times New Roman 12pt."""
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")  # 12 pt = 24 half-points
    rPr.append(sz)
    return rPr


def _make_t(text: str, tag: str = "w:t") -> OxmlElement:
    """Create a <w:t> or <w:delText> element."""
    t = OxmlElement(tag)
    t.text = text
    if text and (text[0] == " " or text[-1] == " "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return t


def _make_run(text: str, rpr: OxmlElement, deleted: bool = False) -> OxmlElement:
    r = OxmlElement("w:r")
    r.append(deepcopy(rpr))
    r.append(_make_t(text, "w:delText" if deleted else "w:t"))
    return r


def make_ins_element(text: str, rpr: OxmlElement) -> OxmlElement:
    """Wrap text in <w:ins>."""
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), _next_id())
    ins.set(qn("w:author"), AUTHOR)
    ins.set(qn("w:date"), DATE)
    ins.append(_make_run(text, rpr))
    return ins


def make_del_element(text: str, rpr: OxmlElement) -> OxmlElement:
    """Wrap text in <w:del>."""
    d = OxmlElement("w:del")
    d.set(qn("w:id"), _next_id())
    d.set(qn("w:author"), AUTHOR)
    d.set(qn("w:date"), DATE)
    d.append(_make_run(text, rpr, deleted=True))
    return d


# ── Character-level PDF extraction ──────────────────────────────────────────

def _group_chars_into_lines(chars, tolerance=3.0):
    """Group characters into lines based on vertical position."""
    if not chars:
        return []

    lines = []
    current_line = [chars[0]]
    current_top = chars[0]["top"]

    for c in chars[1:]:
        if abs(c["top"] - current_top) <= tolerance:
            current_line.append(c)
        else:
            lines.append(current_line)
            current_line = [c]
            current_top = c["top"]
    lines.append(current_line)

    # Sort chars within each line by x position
    for line in lines:
        line.sort(key=lambda c: c["x0"])

    return lines


def _clean_deleted_text(text):
    """Clean up struck-through text by replacing dashes with spaces."""
    # Replace dash characters with spaces
    cleaned = re.sub(r"[\-\u2013\u2014]", " ", text)
    # Collapse multiple spaces
    cleaned = re.sub(r"  +", " ", cleaned)
    return cleaned.strip()


def _is_all_caps_word(word):
    """Check if a word is ALL CAPS (indicating inserted text)."""
    stripped = word.strip(".,;:()'\"[]{}!?")
    if not stripped:
        return False
    # Must have at least 2 alpha chars
    alpha = re.sub(r"[^A-Za-z]", "", stripped)
    if len(alpha) < 2:
        return False
    # Check if all alpha chars are uppercase
    if alpha != alpha.upper():
        return False
    # Exclude dotted abbreviations like U.S.C.
    if ABBREV_RE.match(stripped):
        return False
    # Exclude legal citations like 24-34-301
    if CITATION_RE.match(stripped):
        return False
    return True


def _segment_non_deleted_text(text):
    """
    Segment non-deleted text into (text, 'insert'|'normal') segments
    based on ALL CAPS detection.
    """
    tokens = re.split(r"(\s+)", text)
    segments = []

    for tok in tokens:
        if not tok:
            continue

        if re.match(r"^\s+$", tok):
            # Whitespace: attach to preceding segment
            if segments:
                segments[-1] = (segments[-1][0] + tok, segments[-1][1])
            else:
                segments.append((tok, "normal"))
            continue

        cls = "insert" if _is_all_caps_word(tok) else "normal"

        if segments and segments[-1][1] == cls:
            segments[-1] = (segments[-1][0] + tok, cls)
        else:
            segments.append((tok, cls))

    return segments


def _segment_line_chars(line_chars):
    """
    Segment a line of characters into [(text, class), ...] based on PDF tags.

    - tag='ocr_q' → deleted (struck-through text)
    - other tags with ALL CAPS → inserted
    - other tags with mixed case → normal
    """
    if not line_chars:
        return []

    # First, group consecutive chars by whether they are 'ocr_q' tagged
    raw_segments = []
    cur_text = line_chars[0]["text"]
    cur_is_deleted = line_chars[0].get("tag") == "ocr_q"

    for c in line_chars[1:]:
        is_del = c.get("tag") == "ocr_q"
        if is_del == cur_is_deleted:
            cur_text += c["text"]
        else:
            raw_segments.append((cur_text, cur_is_deleted))
            cur_text = c["text"]
            cur_is_deleted = is_del
    raw_segments.append((cur_text, cur_is_deleted))

    # Now build final segments with classification
    segments = []
    for text, is_deleted in raw_segments:
        if is_deleted:
            cleaned = _clean_deleted_text(text)
            if cleaned:
                segments.append((cleaned, "delete"))
        else:
            # Classify words as insert (ALL CAPS) or normal
            sub_segments = _segment_non_deleted_text(text)
            segments.extend(sub_segments)

    # Merge adjacent segments of the same class
    merged = []
    for text, cls in segments:
        if merged and merged[-1][1] == cls:
            merged[-1] = (merged[-1][0] + text, cls)
        else:
            merged.append((text, cls))

    return merged


def extract_segmented_lines(pdf_path):
    """
    Extract text from PDF as segmented lines.

    Returns a list of lines, where each line is a list of (text, class) tuples.
    class is one of: 'normal', 'insert', 'delete'
    Empty lines are represented as [].
    """
    all_lines = []

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            chars = page.chars
            if not chars:
                continue

            char_lines = _group_chars_into_lines(chars)

            for line_chars in char_lines:
                # Build raw text for noise filtering
                raw_text = "".join(c["text"] for c in line_chars).strip()

                if not raw_text:
                    all_lines.append([])
                    continue

                if NOISE_RE.search(raw_text):
                    continue

                segments = _segment_line_chars(line_chars)
                if segments:
                    all_lines.append(segments)
                else:
                    all_lines.append([])

    return all_lines


def _get_line_text(segments):
    """Get plain text from a segments list."""
    return "".join(text for text, cls in segments)


def _is_paragraph_start(segments):
    """Check if a line starts a new paragraph."""
    text = _get_line_text(segments).strip()
    if not text:
        return True
    # Subsection markers: (a), (b), (1), (2), (I), (II), (III), (5.4), etc.
    if re.match(r"^\([a-zA-Z0-9.]+\)", text):
        return True
    # Section headings: "SECTION 1.", "24-34-301."
    if re.match(r"^SECTION\s+\d+", text):
        return True
    # Statute citation at start of line: "24-34-301."
    if re.match(r"^\d{1,3}-\d{1,5}-\d{1,5}", text):
        return True
    # Bill title lines (all caps, first page)
    if re.match(r"^(HOUSE|SENATE)\s+BILL", text):
        return True
    if re.match(r"^BY\s+(REPRESENTATIVE|SENATOR)", text):
        return True
    if re.match(r"^also\s+SENATOR", text):
        return True
    if re.match(r"^CONCERNING\s", text):
        return True
    if re.match(r"^AND,\s+IN\s+CONNECTION", text):
        return True
    if re.match(r"^Be it enacted", text):
        return True
    return False


def _merge_segment_lists(a, b):
    """Merge two segment lists, joining with a space and merging adjacent same-class segments."""
    if not a:
        return b
    if not b:
        return a

    result = list(a)
    # Add a space to connect lines
    last_cls = result[-1][1]
    first_cls = b[0][1]

    if last_cls == first_cls:
        result[-1] = (result[-1][0] + " " + b[0][0], last_cls)
        result.extend(b[1:])
    else:
        result[-1] = (result[-1][0] + " ", last_cls)
        result.extend(b)

    return result


def join_lines_into_paragraphs(lines):
    """
    Join PDF lines into logical paragraphs.
    A new paragraph starts at blank lines or lines that begin with
    section/subsection markers.
    """
    paragraphs = []
    current = []

    for line in lines:
        is_blank = not line

        if is_blank:
            if current:
                paragraphs.append(current)
                current = []
            paragraphs.append([])  # blank paragraph
            continue

        if _is_paragraph_start(line) and current:
            paragraphs.append(current)
            current = []

        current = _merge_segment_lists(current, line)

    if current:
        paragraphs.append(current)

    # Collapse consecutive blanks and strip leading/trailing
    result = []
    prev_blank = False
    for para in paragraphs:
        is_blank = not para
        if is_blank and prev_blank:
            continue
        result.append(para)
        prev_blank = is_blank
    while result and not result[0]:
        result.pop(0)
    while result and not result[-1]:
        result.pop()
    return result


# ── Word document construction ────────────────────────────────────────────────

def add_paragraph(doc, segments):
    """Append a paragraph containing tracked-change runs for the given segments."""
    para = doc.add_paragraph()
    p_xml = para._p
    rpr = _base_rpr()

    for text, cls in segments:
        if not text:
            continue
        if cls == "normal":
            p_xml.append(_make_run(text, rpr))
        elif cls == "insert":
            p_xml.append(make_ins_element(text, rpr))
        elif cls == "delete":
            p_xml.append(make_del_element(text, rpr))


def build_docx(lines, output_path):
    """Build the Word document with tracked changes from extracted lines."""
    doc = Document()

    # Default paragraph style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # US Letter page size with 1" margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    for segments in lines:
        if not segments:
            doc.add_paragraph()
        else:
            add_paragraph(doc, segments)

    doc.save(output_path)
    print(f"Saved: {output_path}")


# ── Entry point ───────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description=(
            "Convert a Colorado General Assembly bill PDF to a Word document "
            "with tracked changes. Struck-through (dashed) text becomes tracked "
            "deletions; ALL CAPS text becomes tracked insertions. "
            "Author: 'Colorado General Assembly'."
        )
    )
    parser.add_argument("input_pdf", help="Path to input PDF (Colorado GA bill format)")
    parser.add_argument("output_docx", help="Path to output .docx file")
    args = parser.parse_args()

    if not Path(args.input_pdf).exists():
        print(f"Error: File not found: {args.input_pdf}", file=sys.stderr)
        sys.exit(1)

    print(f"Reading:  {args.input_pdf}")
    segmented_lines = extract_segmented_lines(args.input_pdf)
    paragraphs = join_lines_into_paragraphs(segmented_lines)
    print(f"Paragraphs: {len(paragraphs)} (after joining lines)")

    print("Building Word document with tracked changes...")
    build_docx(paragraphs, args.output_docx)


if __name__ == "__main__":
    main()
